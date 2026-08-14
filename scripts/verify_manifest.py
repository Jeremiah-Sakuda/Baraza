#!/usr/bin/env python3
"""``make verify-manifest`` — prove the planted problems are still planted.

Two questions, deliberately never merged into one number:

* **plant** — is the landmine actually in the artifacts on disk? Answerable from
  ``fixtures/corpus/`` alone, on a machine with nothing installed.
* **behaviour** — did the system do the expected thing about it? Answerable only
  from an event log, and reported separately with its own miss list.

``found N of N planted problems`` counts plants. Merging the two counts would
let a corpus that contains every trap stand in for a system that handles them,
which is the more comfortable of the two claims and the less true one.

**The misses are the output.** The miss list prints unconditionally, including
when it is empty, so "no misses" is something this script said rather than a
section that failed to render. A verifier that only reported its successes is
the specific failure mode this target exists to prevent — it is worth stating
plainly, because that verifier is easier to write, prints green more often, and
is worthless.

**The manifest and the probes are cross-checked.** ``fixtures/MANIFEST.md`` owns
the prose, this file owns the probes, and the run fails if the two ID sets
differ. Neither a landmine without a check nor a check without a landmine can
survive a run.

Exit codes, matching ``scripts/compliance.py``:

===  ============================================================
 0   every plant present; every behaviour probe that ran, passed
 1   a plant is missing, or a behaviour probe ran and failed
 2   every plant present, but no event log exists at all
===  ============================================================

Exit 2 is not a pass. It is the honest state of the repository between
``make corpus`` and the first ingest.

Every format is read through ``baraza.ingest.readers`` when its parser is
installed, and through a stdlib fallback in this file when it is not. Which path
was used is printed, because a probe that silently degraded is a probe that
stopped covering the thing it names.
"""

from __future__ import annotations

import argparse
import json
import os
import re
import sys
import zipfile
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any, Callable, Dict, List, Optional, Sequence, Tuple
from xml.etree import ElementTree

REPO = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(REPO / "src"))

from baraza.schema.temporal import (  # noqa: E402
    TemporalError,
    intervals_overlap,
    to_epoch_millis,
)

CORPUS = REPO / "fixtures" / "corpus"
INDEX = CORPUS / "corpus-index.json"
MANIFEST = REPO / "fixtures" / "MANIFEST.md"
GOLD = REPO / "fixtures" / "entities-gold.json"

# Where an event log might be. Ordered by specificity; the first that exists
# wins and the script says which one it used.
LOG_CANDIDATES = ("out/events.jsonl", "fixtures/golden-log.jsonl")

CORPUS_WINDOW = (
    to_epoch_millis("2016-01-01T00:00:00Z"),
    to_epoch_millis("2027-01-01T00:00:00Z"),
)

# The two halves of the L-01 pair, written once so the manifest, this script and
# the fold-stability test all cite the same literals.
TRAP_OFFSET_ISO = "2026-05-01T20:00:00-05:00"
TRAP_UTC_ISO = "2026-05-02T00:00:00Z"

OCR_THRESHOLD = 0.90

DISTINGUISHING = re.compile(
    r"(?i)\b(assistant|deputy|vice|co|interim|acting|former|outgoing|incoming|"
    r"junior|senior|elect)\b"
)


def rel(path: Path) -> str:
    """Repo-relative when possible; absolute otherwise.

    A log or registry passed with --log may live outside the tree, and a path
    formatter that assumed otherwise would crash the run at the report line —
    after every check had already passed.
    """
    try:
        return str(Path(path).relative_to(REPO))
    except ValueError:
        return str(path)


class Unavailable(Exception):
    """A probe could not run. Never the same thing as a probe that failed."""


@dataclass(frozen=True)
class Result:
    ok: bool
    detail: str


def hit(detail: str) -> Result:
    return Result(True, detail)


def miss(detail: str) -> Result:
    return Result(False, detail)


# ------------------------------------------------------- stdlib format readers
#
# Used when a format's third-party parser is absent. They reproduce the locator
# grammar of `baraza.ingest.readers` exactly, because a probe that addressed
# cells differently from the pipeline would verify a different document.

_W = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
_S = "{http://schemas.openxmlformats.org/spreadsheetml/2006/main}"


def _stdlib_xlsx_cells(path: Path) -> Dict[str, str]:
    with zipfile.ZipFile(path) as archive:
        shared: List[str] = []
        if "xl/sharedStrings.xml" in archive.namelist():
            root = ElementTree.fromstring(archive.read("xl/sharedStrings.xml"))
            for si in root.findall(f"{_S}si"):
                shared.append("".join(t.text or "" for t in si.iter(f"{_S}t")))

        workbook = ElementTree.fromstring(archive.read("xl/workbook.xml"))
        names = [s.get("name", "") for s in workbook.iter(f"{_S}sheet")]

        cells: Dict[str, str] = {}
        for position, name in enumerate(names, start=1):
            member = f"xl/worksheets/sheet{position}.xml"
            if member not in archive.namelist():
                continue
            sheet = ElementTree.fromstring(archive.read(member))
            for cell in sheet.iter(f"{_S}c"):
                ref = cell.get("r")
                value = cell.find(f"{_S}v")
                if ref is None or value is None or value.text is None:
                    continue
                text = (
                    shared[int(value.text)]
                    if cell.get("t") == "s"
                    else value.text
                )
                cells[f"{name}!{ref}"] = text
        return cells


def _stdlib_xlsx_rows(path: Path) -> Dict[str, str]:
    """Row context per cell, mirroring ``read_xlsx``'s ``[row: ...]`` suffix."""
    cells = _stdlib_xlsx_cells(path)
    rows: Dict[Tuple[str, str], List[Tuple[int, str]]] = {}
    for ref, value in cells.items():
        sheet, _, coordinate = ref.partition("!")
        match = re.match(r"^([A-Z]+)(\d+)$", coordinate)
        if not match:
            continue
        column = 0
        for char in match.group(1):
            column = column * 26 + (ord(char) - ord("A") + 1)
        rows.setdefault((sheet, match.group(2)), []).append((column, value))
    context: Dict[str, str] = {}
    for ref in cells:
        sheet, _, coordinate = ref.partition("!")
        match = re.match(r"^([A-Z]+)(\d+)$", coordinate)
        if not match:
            continue
        joined = " | ".join(
            v for _, v in sorted(rows[(sheet, match.group(2))])
        )
        context[ref] = joined
    return context


def _stdlib_docx_units(path: Path) -> Dict[str, str]:
    with zipfile.ZipFile(path) as archive:
        root = ElementTree.fromstring(archive.read("word/document.xml"))
    body = root.find(f"{_W}body")
    units: Dict[str, str] = {}
    ordinal = 0
    table_index = 0
    if body is None:
        return units
    for element in body:
        if element.tag == f"{_W}p":
            text = "".join(t.text or "" for t in element.iter(f"{_W}t")).strip()
            if not text:
                continue
            ordinal += 1
            units[f"¶{ordinal}"] = text
        elif element.tag == f"{_W}tbl":
            table_index += 1
            for row_index, row in enumerate(element.findall(f"{_W}tr"), start=1):
                cells = []
                for cell in row.findall(f"{_W}tc"):
                    joined = "".join(
                        t.text or "" for t in cell.iter(f"{_W}t")
                    ).strip()
                    if joined:
                        cells.append(joined)
                if cells:
                    units[f"tbl{table_index}:r{row_index}"] = " | ".join(cells)
    return units


def _stdlib_pdf_pages(path: Path) -> List[str]:
    """Recover the text layer from uncompressed content streams.

    The generator writes streams uncompressed precisely so this works without a
    PDF library; see ``scripts/generate_corpus.py``.
    """
    blob = path.read_bytes().decode("latin-1")
    pages: List[str] = []
    for stream in re.findall(r"stream\n(.*?)endstream", blob, re.DOTALL):
        shown = re.findall(r"\((.*?)\) Tj", stream, re.DOTALL)
        pages.append(
            "\n".join(s.replace(r"\(", "(").replace(r"\)", ")") for s in shown)
        )
    return pages


# ---------------------------------------------------------------- corpus view


@dataclass
class CorpusView:
    """Everything the plant probes read, loaded once and labelled by path."""

    index: Dict[str, Any]
    groupme: Dict[str, Any]
    interview: Dict[str, Any]
    gold: Optional[Dict[str, Any]]
    reader_path: Dict[str, str] = field(default_factory=dict)
    _cells: Dict[str, Dict[str, str]] = field(default_factory=dict)
    _rows: Dict[str, Dict[str, str]] = field(default_factory=dict)
    _docx: Dict[str, Dict[str, str]] = field(default_factory=dict)
    _pdf_units: Dict[str, List[Tuple[str, str, float]]] = field(default_factory=dict)
    _md: Dict[str, str] = field(default_factory=dict)

    @staticmethod
    def load() -> "CorpusView":
        if not INDEX.exists():
            raise Unavailable(
                f"{rel(INDEX)} is missing. Run `make corpus` first."
            )
        index = json.loads(INDEX.read_text(encoding="utf-8"))
        view = CorpusView(
            index=index,
            groupme=json.loads(
                (CORPUS / "chat" / "groupme-meridian-officers.json").read_text(
                    encoding="utf-8"
                )
            ),
            interview=json.loads(
                (
                    CORPUS
                    / "interviews"
                    / "prior-exit-interview-2026-05.json"
                ).read_text(encoding="utf-8")
            ),
            gold=(
                json.loads(GOLD.read_text(encoding="utf-8"))
                if GOLD.exists()
                else None
            ),
        )
        return view

    def path_of(self, source_id: str) -> Path:
        for record in self.index["sources"] + self.index["deferred_sources"]:
            if record["source_id"] == source_id:
                return REPO / record["path"]
        raise Unavailable(f"corpus-index.json does not list source {source_id!r}")

    # -- xlsx ---------------------------------------------------------------

    def cell(self, source_id: str, ref: str) -> str:
        self._load_xlsx(source_id)
        return self._cells[source_id].get(ref, "")

    def row_context(self, source_id: str, ref: str) -> str:
        self._load_xlsx(source_id)
        return self._rows[source_id].get(ref, "")

    def _load_xlsx(self, source_id: str) -> None:
        if source_id in self._cells:
            return
        path = self.path_of(source_id)
        try:
            import openpyxl  # noqa: PLC0415

            workbook = openpyxl.load_workbook(path, data_only=True)
            cells: Dict[str, str] = {}
            for sheet in workbook.worksheets:
                for row in sheet.iter_rows():
                    for cell in row:
                        if cell.value is not None:
                            cells[f"{sheet.title}!{cell.coordinate}"] = str(cell.value)
            workbook.close()
            self.reader_path[source_id] = "openpyxl"
        except ImportError:
            cells = _stdlib_xlsx_cells(path)
            self.reader_path[source_id] = "stdlib zip+xml fallback"
        self._cells[source_id] = cells
        # Row context always comes from the stdlib pass: it reproduces the
        # "[row: ...]" suffix read_xlsx attaches to every cell unit, which is
        # what makes the L-09 decoy visible to a human reading the ledger.
        self._rows[source_id] = _stdlib_xlsx_rows(path)

    # -- docx ---------------------------------------------------------------

    def docx(self, source_id: str) -> Dict[str, str]:
        if source_id not in self._docx:
            path = self.path_of(source_id)
            try:
                import docx  # noqa: PLC0415

                document = docx.Document(str(path))
                units: Dict[str, str] = {}
                ordinal = 0
                for paragraph in document.paragraphs:
                    text = paragraph.text.strip()
                    if text:
                        ordinal += 1
                        units[f"¶{ordinal}"] = text
                for table_index, table in enumerate(document.tables, start=1):
                    for row_index, row in enumerate(table.rows, start=1):
                        cells = [c.text.strip() for c in row.cells if c.text.strip()]
                        if cells:
                            units[f"tbl{table_index}:r{row_index}"] = " | ".join(cells)
                self.reader_path[source_id] = "python-docx"
            except ImportError:
                units = _stdlib_docx_units(path)
                self.reader_path[source_id] = "stdlib zip+xml fallback"
            self._docx[source_id] = units
        return self._docx[source_id]

    # -- pdf ----------------------------------------------------------------

    def pdf_units(self, source_id: str) -> List[Tuple[str, str, float]]:
        """``(locator, text, confidence)`` for every unit of a scanned source."""
        if source_id not in self._pdf_units:
            path = self.path_of(source_id)
            units: List[Tuple[str, str, float]] = []
            try:
                from baraza.ingest.readers import read_source  # noqa: PLC0415

                source = read_source(
                    path, source_id=source_id, observed_at="2024-03-12T00:00:00Z"
                )
                units = [
                    (u.locator, u.text, u.confidence) for u in source.units.values()
                ]
                self.reader_path[source_id] = "baraza.ingest.readers"
            except Exception:  # MissingReaderDependency or an absent parser
                from baraza.ingest.readers import _ocr_confidence  # noqa: PLC0415

                for page_number, text in enumerate(
                    _stdlib_pdf_pages(path), start=1
                ):
                    units.append(
                        (f"p.{page_number} ¶1", text, _ocr_confidence(text))
                    )
                self.reader_path[source_id] = "stdlib pdf text-layer fallback"
            self._pdf_units[source_id] = units
        return self._pdf_units[source_id]

    def pdf_text(self, source_id: str) -> str:
        return "\n".join(text for _, text, _ in self.pdf_units(source_id))

    # -- md -----------------------------------------------------------------

    def md_text(self, source_id: str) -> str:
        if source_id not in self._md:
            self._md[source_id] = self.path_of(source_id).read_text(encoding="utf-8")
            self.reader_path[source_id] = "utf-8 text"
        return self._md[source_id]

    # -- groupme ------------------------------------------------------------

    def message(self, locator: str) -> Optional[Dict[str, Any]]:
        wanted = locator.removeprefix("msg:")
        for entry in self.groupme["messages"]:
            if str(entry["created_at"]) == wanted:
                return entry
        return None

    def segment(self, segment_id: str) -> Optional[Dict[str, Any]]:
        for entry in self.groupme["segments"]:
            if entry["segment_id"] == segment_id:
                return entry
        return None


# ------------------------------------------------------------------ log view


@dataclass
class LogView:
    """The folded event log, if there is one."""

    path: Path
    claims: List[Any]
    contradictions: List[Any]
    aliases: Dict[str, str]

    @staticmethod
    def load(explicit: Optional[str]) -> "LogView":
        from baraza.fold.graph import fold  # noqa: PLC0415
        from baraza.fold.store import JsonlEventStore  # noqa: PLC0415

        candidates = (
            [Path(explicit)]
            if explicit
            else (
                [Path(os.environ["BARAZA_EVENT_LOG"])]
                if os.environ.get("BARAZA_EVENT_LOG")
                else [REPO / c for c in LOG_CANDIDATES]
            )
        )
        for candidate in candidates:
            if candidate.exists() and candidate.stat().st_size > 0:
                state = fold(JsonlEventStore(candidate).read_all())
                return LogView(
                    path=candidate,
                    claims=list(state.claims.values()),
                    contradictions=list(state.contradictions.values()),
                    aliases=dict(state.aliases),
                )
        looked = ", ".join(str(c) for c in candidates)
        raise Unavailable(
            f"no non-empty event log found (looked at: {looked}). "
            "Behaviour cannot be observed until an ingest run has happened; "
            "try `make demo-agenda` first."
        )

    def by_source(self, source_id: str) -> List[Any]:
        return [c for c in self.claims if c.anchor.source_id == source_id]

    def anchored_at(self, source_id: str, locator: str) -> List[Any]:
        return [
            c
            for c in self.claims
            if c.anchor.source_id == source_id and c.anchor.locator == locator
        ]

    def quote_of(self, claim: Any) -> str:
        """Read a claim's quote through the audience predicate, as every read
        path in the system must. ``OWNER`` is the audit clearance: this script
        checks grounding across the whole log, including private claims, and the
        audience is passed explicitly rather than inferred."""
        from baraza.schema.visibility import Audience  # noqa: PLC0415

        return claim.quote_for(Audience.OWNER) or ""

    def sides(self, contradiction: Any) -> List[Any]:
        index = {c.claim_id: c for c in self.claims}
        return [index[cid] for cid in contradiction.claim_ids if cid in index]

    def spans(self, contradiction: Any, left: str, right: str) -> bool:
        sources = {c.anchor.source_id for c in self.sides(contradiction)}
        return left in sources and right in sources


# ------------------------------------------------------------------- landmines


@dataclass(frozen=True)
class Landmine:
    landmine_id: str
    title: str
    plant: Callable[[CorpusView], Result]
    behaviour: Optional[Callable[[CorpusView, LogView], Result]]
    delegated_to: str = ""
    """Where the behaviour is verified when it is not verified here. Empty
    unless ``behaviour`` is None; a landmine may not opt out silently."""


def _contains(haystack: str, needle: str) -> bool:
    return re.sub(r"\s+", " ", needle).lower() in re.sub(
        r"\s+", " ", haystack
    ).lower()


def _find_unit(view: CorpusView, source_id: str, needle: str) -> Optional[Tuple[str, str, float]]:
    """Locate planted PDF text by content, never by a fixed ¶ ordinal.

    pypdf and pdfplumber disagree about how many units a page of this document
    has (33 vs 7, measured). The page component is stable; the paragraph
    ordinal is not, so probes search.
    """
    for locator, text, confidence in view.pdf_units(source_id):
        if _contains(text, needle):
            return locator, text, confidence
    return None


def _normalize_amount(text: str) -> Optional[float]:
    cleaned = re.sub(r"[^0-9.\-]", "", (text or "").replace(",", ""))
    if not cleaned or cleaned in ("-", ".", "-."):
        return None
    try:
        return float(cleaned)
    except ValueError:
        return None


def build_landmines() -> List[Landmine]:
    L: List[Landmine] = []

    # -- L-01 ---------------------------------------------------------------
    def l01_plant(view: CorpusView) -> Result:
        segment = view.segment("seg-2026-may")
        if segment is None:
            return miss("groupme export has no segment seg-2026-may")
        if segment.get("segment_started_at_iso") != TRAP_OFFSET_ISO:
            return miss(
                f"seg-2026-may starts at {segment.get('segment_started_at_iso')!r}, "
                f"expected {TRAP_OFFSET_ISO!r}"
            )
        if segment.get("tz_offset") != "-05:00":
            return miss(f"seg-2026-may tz_offset is {segment.get('tz_offset')!r}")
        turns = {t["turn_id"]: t for t in view.interview["turns"]}
        if turns.get("t-2", {}).get("ts") != TRAP_UTC_ISO:
            return miss(
                f"interview turn t-2 ts is {turns.get('t-2', {}).get('ts')!r}, "
                f"expected {TRAP_UTC_ISO!r}"
            )
        text_order = TRAP_OFFSET_ISO < TRAP_UTC_ISO
        instant_order = to_epoch_millis(TRAP_OFFSET_ISO) < to_epoch_millis(
            TRAP_UTC_ISO
        )
        if not (text_order and not instant_order):
            return miss(
                "the planted pair no longer diverges: string order "
                f"{text_order}, instant order {instant_order}. A trap that "
                "agrees with itself tests nothing."
            )
        return hit(
            f"string order True, instant order False "
            f"({to_epoch_millis(TRAP_OFFSET_ISO)} vs "
            f"{to_epoch_millis(TRAP_UTC_ISO)} ms)"
        )

    def l01_behaviour(view: CorpusView, log: LogView) -> Result:
        low, high = CORPUS_WINDOW
        stray = [
            c for c in log.claims if not (low <= c.observed_at < high)
        ]
        if stray:
            worst = stray[0]
            return miss(
                f"{len(stray)} claim(s) observed outside the corpus window; "
                f"e.g. {worst.claim_id} at {worst.observed_at} ms "
                f"({worst.anchor.key()})"
            )
        return hit(f"all {len(log.claims)} claims inside the corpus window")

    L.append(Landmine("L-01", "mixed-UTC-offset trap", l01_plant, l01_behaviour))

    # -- L-02 ---------------------------------------------------------------
    def l02_plant(view: CorpusView) -> Result:
        a = view.docx("minutes-2023-09-12").get("¶6", "")
        b = view.docx("minutes-2024-09-10").get("¶6", "")
        if "1 July 2023 through 30 June 2024" not in a:
            return miss(f"minutes-2023-09-12 ¶6 does not declare FY24: {a[:70]!r}")
        if "1 July 2024 through 30 June 2025" not in b:
            return miss(f"minutes-2024-09-10 ¶6 does not declare FY25: {b[:70]!r}")
        overlap = intervals_overlap(
            to_epoch_millis("2023-07-01"),
            to_epoch_millis("2024-06-30"),
            to_epoch_millis("2024-07-01"),
            to_epoch_millis("2025-06-30"),
        )
        if overlap:
            return miss(
                "the two declared terms overlap; the false positive would be a "
                "true positive and the fixture proves nothing"
            )
        return hit("FY24 and FY25 terms declared and disjoint on epoch values")

    def l02_behaviour(view: CorpusView, log: LogView) -> Result:
        offenders = []
        for contradiction in log.contradictions:
            sides = log.sides(contradiction)
            if len(sides) != 2:
                continue
            left, right = sides
            if not intervals_overlap(
                left.valid_from, left.valid_until, right.valid_from, right.valid_until
            ):
                offenders.append(contradiction.contradiction_id)
        if offenders:
            return miss(
                f"{len(offenders)} contradiction(s) hold non-overlapping "
                f"intervals — the temporal gate did not fire: {offenders[:3]}"
            )
        return hit(
            f"no contradiction spans disjoint intervals "
            f"({len(log.contradictions)} on the ledger)"
        )

    L.append(
        Landmine("L-02", "FY-pair false positive", l02_plant, l02_behaviour)
    )

    # -- L-03 ---------------------------------------------------------------
    def l03_plant(view: CorpusView) -> Result:
        found = _find_unit(
            view, "constitution", "two hundred fifty dollars"
        )
        if found is None:
            return miss("constitution has no $250 two-signature clause")
        minutes = view.docx("minutes-2021-11-09")
        motion = [t for t in minutes.values() if "single signature" in t]
        if not motion:
            return miss("minutes-2021-11-09 records no single-signature motion")
        if "No amendment to the Constitution was proposed" not in " ".join(
            minutes.values()
        ):
            return miss(
                "minutes-2021-11-09 no longer records that no amendment was made "
                "— without that line the contradiction is arguable"
            )
        return hit(f"constitution {found[0]} vs minutes-2021-11-09 motion")

    def l03_behaviour(view: CorpusView, log: LogView) -> Result:
        matched = [
            c
            for c in log.contradictions
            if log.spans(c, "constitution", "minutes-2021-11-09")
        ]
        if not matched:
            return miss(
                "no contradiction spans constitution and minutes-2021-11-09; "
                f"ledger holds {len(log.contradictions)} row(s)"
            )
        return hit(f"{len(matched)} contradiction(s), e.g. {matched[0].contradiction_id}")

    L.append(
        Landmine("L-03", "signing authority: constitution vs minutes", l03_plant, l03_behaviour)
    )

    # -- L-04 ---------------------------------------------------------------
    def l04_plant(view: CorpusView) -> Result:
        message = view.message("msg:1700001660")
        if message is None or "dues are still 25" not in message["text"]:
            return miss("chat message msg:1700001660 no longer states $25 dues")
        if view.cell("budget-workbook", "Sheet1!B3") != "40":
            return miss(
                f"Sheet1!B3 is {view.cell('budget-workbook', 'Sheet1!B3')!r}, "
                "expected '40'"
            )
        context = view.row_context("budget-workbook", "Sheet1!B3")
        if "dues income" not in context:
            return miss(f"Sheet1!B3's row is not the dues row: {context!r}")
        return hit("chat says 25, Sheet1!B3 says 40 in the dues-income row")

    def l04_behaviour(view: CorpusView, log: LogView) -> Result:
        matched = [
            c for c in log.contradictions if log.spans(c, "gm-officers", "budget-workbook")
        ]
        if not matched:
            return miss(
                "no contradiction spans gm-officers and budget-workbook"
            )
        return hit(f"{len(matched)} contradiction(s), e.g. {matched[0].contradiction_id}")

    L.append(Landmine("L-04", "dues amount: chat vs spreadsheet", l04_plant, l04_behaviour))

    # -- L-05 ---------------------------------------------------------------
    def l05_plant(view: CorpusView) -> Result:
        found = _find_unit(
            view, "constitution", "no independent authority to disburse funds"
        )
        if found is None:
            return miss("constitution no longer bars the Assistant Treasurer")
        text = view.docx("minutes-2023-09-12").get("¶7", "")
        if "does not sign on the account" not in text:
            return miss(f"minutes-2023-09-12 ¶7 lost the distinction: {text[:70]!r}")
        if view.gold is None:
            return miss("fixtures/entities-gold.json is missing")
        distinct = {
            tuple(sorted(entry["pair"])) for entry in view.gold.get("distinct", [])
        }
        pair = tuple(sorted(("ent:treasurer", "ent:assistant-treasurer")))
        if pair not in distinct:
            return miss("gold file does not list treasurer/assistant-treasurer as distinct")
        return hit(f"constitution {found[0]}, minutes ¶7, gold `distinct` entry")

    def l05_behaviour(view: CorpusView, log: LogView) -> Result:
        bad = [
            (alias, canonical)
            for alias, canonical in log.aliases.items()
            if "treasurer" in alias.lower()
            and "treasurer" in canonical.lower()
            and bool(DISTINGUISHING.search(alias)) != bool(
                DISTINGUISHING.search(canonical)
            )
        ]
        if bad:
            return miss(f"an alias edge merged distinct offices: {bad}")
        return hit(f"no office-modifier alias edge ({len(log.aliases)} edges total)")

    L.append(
        Landmine("L-05", "Treasurer vs Assistant Treasurer must not merge", l05_plant, l05_behaviour)
    )

    # -- L-06 ---------------------------------------------------------------
    def l06_plant(view: CorpusView) -> Result:
        units = view.pdf_units("constitution")
        if not units:
            return miss("constitution produced no units")
        degraded = [u for u in units if u[2] < OCR_THRESHOLD]
        if len(degraded) != 1:
            return miss(
                f"{len(degraded)} unit(s) below {OCR_THRESHOLD}, expected exactly "
                f"1; worst is {min(u[2] for u in units):.3f}"
            )
        locator, text, confidence = degraded[0]
        if "$?SO" not in text:
            return miss(
                f"the degraded unit {locator} has lost the illegible-amount "
                f"marker: {text[:70]!r}"
            )
        next_worst = min(
            (u[2] for u in units if u[2] >= OCR_THRESHOLD), default=1.0
        )
        return hit(
            f"{locator} at {confidence:.3f}, next-worst {next_worst:.3f}, "
            "illegible amount present"
        )

    def l06_behaviour(view: CorpusView, log: LogView) -> Result:
        degraded = [u for u in view.pdf_units("constitution") if u[2] < OCR_THRESHOLD]
        if not degraded:
            return miss("no degraded unit to check claims against")
        locator = degraded[0][0]
        anchored = log.anchored_at("constitution", locator)
        unflagged = [
            c
            for c in anchored
            if float(c.extra.get("unit_confidence", 1.0)) >= OCR_THRESHOLD
        ]
        if unflagged:
            return miss(
                f"{len(unflagged)} claim(s) anchored at {locator} carry no "
                f"low-confidence flag, e.g. {unflagged[0].claim_id}"
            )
        asserted = [
            c
            for c in anchored
            if _normalize_amount(c.object_literal or "") is not None
        ]
        if asserted:
            return miss(
                f"{asserted[0].claim_id} asserts the amount "
                f"{asserted[0].object_literal!r} from an illegible region"
            )
        return hit(
            f"{len(anchored)} claim(s) at {locator}, all flagged, none asserting "
            "the illegible amount"
        )

    L.append(Landmine("L-06", "low-OCR-confidence region", l06_plant, l06_behaviour))

    # -- L-07 ---------------------------------------------------------------
    def l07_plant(view: CorpusView) -> Result:
        big = [
            m for m in view.groupme["messages"] if int(m["created_at"]) >= 10_000_000_000
        ]
        if len(big) != 1:
            return miss(f"{len(big)} millisecond timestamps, expected exactly 1")
        instant = to_epoch_millis(big[0]["created_at"])
        low, high = CORPUS_WINDOW
        if not low <= instant < high:
            return miss(
                f"the millisecond message normalizes to {instant} ms, outside "
                "the corpus window"
            )
        return hit(f"msg:{big[0]['created_at']} normalizes to {instant} ms")

    def l07_behaviour(view: CorpusView, log: LogView) -> Result:
        expected = to_epoch_millis("2024-10-15T19:47:00Z")
        anchored = log.anchored_at("gm-officers", "msg:1729021620000")
        drifted = [
            c for c in anchored if abs(c.observed_at - expected) > 86_400_000
        ]
        if drifted:
            return miss(
                f"{drifted[0].claim_id} observed at {drifted[0].observed_at} ms, "
                f"more than a day from {expected} ms — read as seconds"
            )
        return hit(
            f"{len(anchored)} claim(s) at the millisecond message, all within a "
            "day of the true instant"
        )

    L.append(Landmine("L-07", "epoch seconds vs milliseconds", l07_plant, l07_behaviour))

    # -- L-08 ---------------------------------------------------------------
    def l08_plant(view: CorpusView) -> Result:
        text = view.md_text("notes-handover")
        if "2026-04-14T19:30:00" not in text:
            return miss("the offsetless timestamp is gone from the handover note")
        if "2026-04-14T19:30:00Z" in text or "2026-04-14T19:30:00-" in text:
            return miss("the timestamp acquired an offset; it is no longer ambiguous")
        try:
            to_epoch_millis("2026-04-14T19:30:00")
        except TemporalError:
            return hit("present, and to_epoch_millis raises TemporalError on it")
        return miss(
            "to_epoch_millis accepted an offsetless ISO string — BAR-309 forbids "
            "guessing an offset"
        )

    def l08_behaviour(view: CorpusView, log: LogView) -> Result:
        guessed = to_epoch_millis("2026-04-14T19:30:00Z")
        offenders = [
            c
            for c in log.claims
            if guessed in (c.observed_at, c.valid_from, c.valid_until)
        ]
        if offenders:
            return miss(
                f"{offenders[0].claim_id} carries the guessed-UTC reading of an "
                "offsetless local time"
            )
        return hit("no claim carries the guessed-UTC instant")

    L.append(Landmine("L-08", "offsetless ISO timestamp", l08_plant, l08_behaviour))

    # -- L-09 ---------------------------------------------------------------
    def l09_plant(view: CorpusView) -> Result:
        dues = view.cell("budget-workbook", "Sheet1!B3")
        rental = view.cell("budget-workbook", "Sheet1!B7")
        if dues != "40" or rental != "40":
            return miss(f"decoy gone: B3={dues!r}, B7={rental!r}, both should be '40'")
        dues_row = view.row_context("budget-workbook", "Sheet1!B3")
        rental_row = view.row_context("budget-workbook", "Sheet1!B7")
        if "dues income" not in dues_row or "chair rental" not in rental_row:
            return miss(
                f"row labels no longer distinguish the cells: {dues_row!r} / "
                f"{rental_row!r}"
            )
        return hit("B3 and B7 both '40', in the dues-income and chair-rental rows")

    def l09_behaviour(view: CorpusView, log: LogView) -> Result:
        wrong = [
            c
            for c in log.anchored_at("budget-workbook", "Sheet1!B7")
            if "dues" in c.predicate_hint.lower()
        ]
        if wrong:
            return miss(
                f"{wrong[0].claim_id} makes a dues claim from the chair-rental "
                "cell — grounded and wrong"
            )
        return hit("no dues claim cites Sheet1!B7")

    L.append(Landmine("L-09", "headerless-column decoy", l09_plant, l09_behaviour))

    # -- L-10 ---------------------------------------------------------------
    def l10_plant(view: CorpusView) -> Result:
        for locator, needle in (
            ("msg:1722359700", "new handle incoming"),
            ("msg:1724083320", "it is me, sablewick"),
            ("msg:1727212440", "two different people"),
        ):
            message = view.message(locator)
            if message is None or needle not in message["text"]:
                return miss(f"{locator} no longer carries {needle!r}")
        if view.gold is None:
            return miss("fixtures/entities-gold.json is missing")
        same = {tuple(sorted(e["pair"])) for e in view.gold.get("same_as", [])}
        distinct = {tuple(sorted(e["pair"])) for e in view.gold.get("distinct", [])}
        if tuple(sorted(("ent:sablewick", "ent:sable.w"))) not in same:
            return miss("gold file does not label sablewick/sable.w as the same human")
        if tuple(sorted(("ent:sablewick", "ent:sablewood"))) not in distinct:
            return miss("gold file does not label sablewick/sablewood as distinct")
        return hit("three chat messages present; both gold labels in place")

    def l10_behaviour(view: CorpusView, log: LogView) -> Result:
        pair = {"ent:sablewick", "ent:sablewood"}
        for alias, canonical in log.aliases.items():
            if {alias, canonical} == pair:
                return miss(f"alias edge merged two different humans: {alias} -> {canonical}")
        return hit("sablewick and sablewood are not linked")

    L.append(Landmine("L-10", "handle change vs lookalike handle", l10_plant, l10_behaviour))

    # -- L-11 ---------------------------------------------------------------
    def l11_plant(view: CorpusView) -> Result:
        paraphrase = "requires two signatures on all disbursements"
        text = view.docx("minutes-2023-10-10").get("¶7", "")
        if paraphrase not in text:
            return miss(f"minutes-2023-10-10 ¶7 lost the paraphrase: {text[:70]!r}")
        if _contains(view.pdf_text("constitution"), "on all disbursements"):
            return miss(
                "the constitution now contains the paraphrase; the trap only "
                "works while the source and the paraphrase differ"
            )
        return hit("paraphrase in the minutes, absent from the constitution")

    def l11_behaviour(view: CorpusView, log: LogView) -> Result:
        offenders = [
            c
            for c in log.by_source("constitution")
            if "on all disbursements" in log.quote_of(c).lower()
        ]
        if offenders:
            return miss(
                f"{offenders[0].claim_id} attributes the minutes' paraphrase to "
                "the constitution"
            )
        return hit("no constitution-anchored claim carries the paraphrase")

    L.append(Landmine("L-11", "paraphrase attributed to the wrong document", l11_plant, l11_behaviour))

    # -- L-12 ---------------------------------------------------------------
    def l12_plant(view: CorpusView) -> Result:
        for locator, needle in (
            ("msg:1731449280", "signed the venue deposit today, 600"),
            ("msg:1731449700", "are you allowed to do that"),
            ("msg:1731449940", "nobody stopped me"),
        ):
            message = view.message(locator)
            if message is None or needle not in message["text"]:
                return miss(f"{locator} no longer carries {needle!r}")
        return hit("the three authority-creep messages are present")

    def l12_behaviour(view: CorpusView, log: LogView) -> Result:
        locators = ("msg:1731449280", "msg:1731449700", "msg:1731449940")
        anchored = [c for loc in locators for c in log.anchored_at("gm-officers", loc)]
        if not anchored:
            return miss(
                "no claim anchored in the authority-creep thread; the pre-filter "
                "dropped the highest-value finding in the corpus"
            )
        return hit(f"{len(anchored)} claim(s) survived the pre-filter")

    L.append(Landmine("L-12", "authority creep by the Assistant Treasurer", l12_plant, l12_behaviour))

    # -- L-13 ---------------------------------------------------------------
    def l13_plant(view: CorpusView) -> Result:
        bots = [m for m in view.groupme["messages"] if m.get("sender_type") == "bot"]
        if not bots:
            return miss("no message is labelled sender_type=bot")
        return hit(f"{len(bots)} bot post(s), labelled in the export")

    def l13_behaviour(view: CorpusView, log: LogView) -> Result:
        locators = {
            f"msg:{m['created_at']}"
            for m in view.groupme["messages"]
            if m.get("sender_type") == "bot"
        }
        offenders = [
            c for c in log.by_source("gm-officers") if c.anchor.locator in locators
        ]
        if offenders:
            return miss(
                f"{offenders[0].claim_id} treats an automated post as testimony "
                f"({offenders[0].anchor.key()})"
            )
        return hit("no claim is anchored at an automated post")

    L.append(Landmine("L-13", "automated posts are not organic activity", l13_plant, l13_behaviour))

    # -- L-14 ---------------------------------------------------------------
    def l14_plant(view: CorpusView) -> Result:
        if view.cell("budget-workbook", "Sheet1 (2)!B4") != "1,250.00":
            return miss(
                f"Sheet1 (2)!B4 is "
                f"{view.cell('budget-workbook', 'Sheet1 (2)!B4')!r}, expected the "
                "text form '1,250.00'"
            )
        message = view.message("msg:1651605600")
        if message is None or "1250" not in message["text"]:
            return miss("the bare '1250' form is gone from the chat")
        if "$1,250" not in view.md_text("notes-postmortem"):
            return miss("the '$1,250' form is gone from the postmortem note")
        return hit("'1,250.00', '1250' and '$1,250' all present")

    def l14_behaviour(view: CorpusView, log: LogView) -> Result:
        for contradiction in log.contradictions:
            sides = log.sides(contradiction)
            if len(sides) != 2:
                continue
            left = _normalize_amount(sides[0].object_literal or "")
            right = _normalize_amount(sides[1].object_literal or "")
            if left is not None and left == right:
                return miss(
                    f"{contradiction.contradiction_id} treats two spellings of "
                    f"{left} as a disagreement"
                )
        return hit("no contradiction rests on two spellings of one number")

    L.append(Landmine("L-14", "three spellings of one number", l14_plant, l14_behaviour))

    # -- L-15 ---------------------------------------------------------------
    def l15_plant(view: CorpusView) -> Result:
        deferred = {r["source_id"] for r in view.index["deferred_sources"]}
        night1 = {r["source_id"] for r in view.index["sources"]}
        if "minutes-2026-04-14" not in deferred:
            return miss("minutes-2026-04-14 is not listed under deferred_sources")
        if "minutes-2026-04-14" in night1:
            return miss(
                "minutes-2026-04-14 is in the night-1 sources; there is nothing "
                "left to drop and the differential ledger would be empty"
            )
        path = CORPUS / "minutes" / "minutes-2026-04-14.docx"
        if not path.exists():
            return miss(f"{rel(path)} was not generated")
        text = " ".join(_stdlib_docx_units(path).values())
        if "any two officers jointly" not in text:
            return miss("the deferred minutes no longer introduce a fourth position")
        return hit("generated, deferred, and carrying a new signing-authority position")

    L.append(
        Landmine(
            "L-15",
            "deferred artifact drop (BAR-323)",
            l15_plant,
            None,
            delegated_to=(
                "BAR-323 differential choreography — needs two genuinely "
                "separated nightly runs; see docs/GATE.md G3"
            ),
        )
    )

    # -- L-16 ---------------------------------------------------------------
    def l16_plant(view: CorpusView) -> Result:
        listed = [
            r["path"]
            for r in view.index["sources"] + view.index["deferred_sources"]
        ]
        offenders = [p for p in listed if p.endswith("BIBLE.md")]
        if offenders:
            return miss(
                f"the answer key is listed as an ingestible source: {offenders}. "
                "Every downstream number would be measured against a test the "
                "system had been handed the answers to."
            )
        not_sources = [r["path"] for r in view.index.get("not_sources", [])]
        if not any(p.endswith("BIBLE.md") for p in not_sources):
            return miss(
                "BIBLE.md is absent from not_sources; exclusion by omission is "
                "not exclusion by decision"
            )
        return hit("BIBLE.md excluded from sources and named in not_sources")

    def l16_behaviour(view: CorpusView, log: LogView) -> Result:
        offenders = [c for c in log.claims if c.anchor.source_id.lower() in ("bible", "bible.md")]
        if offenders:
            return miss(f"{offenders[0].claim_id} cites the answer key")
        return hit("no claim cites the answer key")

    L.append(Landmine("L-16", "the answer key is never ingested", l16_plant, l16_behaviour))

    # -- L-17 ---------------------------------------------------------------
    def l17_plant(view: CorpusView) -> Result:
        first = view.message("msg:1701368520")
        second = view.message("msg:1701368760")
        if first is None or "equipment fund is at 640" not in first["text"]:
            return miss("the retracted figure is gone from msg:1701368520")
        if second is None or "scratch that" not in second["text"]:
            return miss("the retraction is gone from msg:1701368760")
        gap = int(second["created_at"]) - int(first["created_at"])
        if gap <= 0:
            return miss(f"the retraction does not follow the claim (gap {gap}s)")
        return hit(f"claim and retraction {gap}s apart, in order")

    def l17_behaviour(view: CorpusView, log: LogView) -> Result:
        wrong = log.anchored_at("gm-officers", "msg:1701368520")
        correction = log.anchored_at("gm-officers", "msg:1701368760")
        if wrong and not correction:
            return miss(
                "the retracted figure was extracted but its correction was not; "
                "the ledger would carry a number the corpus disowns"
            )
        return hit(
            f"{len(wrong)} claim(s) on the retracted figure, "
            f"{len(correction)} on the correction"
        )

    L.append(Landmine("L-17", "a figure retracted four minutes later", l17_plant, l17_behaviour))

    # -- L-18 ---------------------------------------------------------------
    def l18_plant(view: CorpusView) -> Result:
        found = _find_unit(
            view, "constitution", "shall not be altered during the year then in progress"
        )
        if found is None:
            return miss("constitution no longer forbids a mid-year dues change")
        text = view.docx("minutes-2023-10-10").get("¶6", "")
        if "effective the spring semester only" not in text:
            return miss(f"minutes-2023-10-10 ¶6 lost the mid-year raise: {text[:70]!r}")
        return hit(f"constitution {found[0]} vs minutes-2023-10-10 ¶6")

    def l18_behaviour(view: CorpusView, log: LogView) -> Result:
        matched = [
            c
            for c in log.contradictions
            if log.spans(c, "constitution", "minutes-2023-10-10")
        ]
        if not matched:
            return miss("no contradiction spans constitution and minutes-2023-10-10")
        return hit(f"{len(matched)} contradiction(s), e.g. {matched[0].contradiction_id}")

    L.append(Landmine("L-18", "dues altered mid-year", l18_plant, l18_behaviour))

    return L


# ------------------------------------------------------------- manifest sync


def manifest_ids() -> List[str]:
    if not MANIFEST.exists():
        raise Unavailable(f"{rel(MANIFEST)} is missing")
    return sorted(
        set(
            re.findall(
                r"^##\s+(L-\d{2})\s+—", MANIFEST.read_text(encoding="utf-8"), re.M
            )
        )
    )


# -------------------------------------------------------------------- runner


def main(argv: Sequence[str]) -> int:
    parser = argparse.ArgumentParser(description="verify the planted problems")
    parser.add_argument("--log", help="path to an event log (JSONL)")
    parser.add_argument(
        "--plants-only",
        action="store_true",
        help="skip the behaviour phase entirely and say so",
    )
    args = parser.parse_args(argv)

    landmines = build_landmines()
    print("make verify-manifest — fixtures/MANIFEST.md")
    print("=" * 78)

    # -- manifest/probe cross-check ----------------------------------------
    try:
        documented = manifest_ids()
    except Unavailable as exc:
        print(f"cannot run: {exc}")
        return 2

    probed = sorted(m.landmine_id for m in landmines)
    if documented != probed:
        print("MANIFEST and probe set diverge — neither can be trusted:")
        for missing in sorted(set(documented) - set(probed)):
            print(f"  {missing}: documented in MANIFEST.md, no probe in this script")
        for extra in sorted(set(probed) - set(documented)):
            print(f"  {extra}: probed by this script, absent from MANIFEST.md")
        return 1
    print(f"  manifest/probe cross-check   ok ({len(probed)} ids)")

    # -- plants -------------------------------------------------------------
    try:
        view = CorpusView.load()
    except Unavailable as exc:
        print(f"cannot run: {exc}")
        return 2
    except FileNotFoundError as exc:
        print(f"cannot run: a corpus artifact is missing ({exc}). Run `make corpus`.")
        return 2

    plant_results: Dict[str, Result] = {}
    for landmine in landmines:
        try:
            plant_results[landmine.landmine_id] = landmine.plant(view)
        except Exception as exc:  # a raising probe is a failing probe
            plant_results[landmine.landmine_id] = miss(
                f"probe raised {type(exc).__name__}: {exc}"
            )

    # -- behaviour ----------------------------------------------------------
    log: Optional[LogView] = None
    log_error = ""
    if not args.plants_only:
        try:
            log = LogView.load(args.log)
        except Unavailable as exc:
            log_error = str(exc)
    else:
        log_error = "--plants-only was passed; the behaviour phase did not run"

    behaviour_results: Dict[str, Result] = {}
    if log is not None:
        for landmine in landmines:
            if landmine.behaviour is None:
                continue
            try:
                behaviour_results[landmine.landmine_id] = landmine.behaviour(view, log)
            except Exception as exc:
                behaviour_results[landmine.landmine_id] = miss(
                    f"probe raised {type(exc).__name__}: {exc}"
                )

    # -- report -------------------------------------------------------------
    print()
    print(f"  {'id':<6} {'plant':<6} {'behaviour':<11} landmine")
    print("  " + "-" * 74)
    for landmine in landmines:
        plant = "ok" if plant_results[landmine.landmine_id].ok else "MISS"
        if landmine.behaviour is None:
            behaviour = "delegated"
        elif landmine.landmine_id in behaviour_results:
            behaviour = "ok" if behaviour_results[landmine.landmine_id].ok else "MISS"
        else:
            behaviour = "not-run"
        print(
            f"  {landmine.landmine_id:<6} {plant:<6} {behaviour:<11} {landmine.title}"
        )

    print()
    print("  readers used:")
    for source_id, path in sorted(view.reader_path.items()):
        print(f"    {source_id:<20} {path}")

    plant_hits = sum(1 for r in plant_results.values() if r.ok)
    total = len(landmines)
    behaviour_total = sum(1 for m in landmines if m.behaviour is not None)
    behaviour_hits = sum(1 for r in behaviour_results.values() if r.ok)

    print()
    print("=" * 78)
    print(f"found {plant_hits} of {total} planted problems")
    if log is not None:
        print(
            f"behaviour observed: {behaviour_hits} of {behaviour_total} probes "
            f"(event log: {rel(log.path)})"
        )
    else:
        print(f"behaviour observed: 0 of {behaviour_total} probes — {log_error}")
    delegated = [m for m in landmines if m.behaviour is None]
    for landmine in delegated:
        print(
            f"behaviour delegated: {landmine.landmine_id} -> {landmine.delegated_to}"
        )
    print("=" * 78)

    # The miss list prints unconditionally, empty or not. This is the whole
    # point of the target: a run that only listed its successes would look
    # identical to a run that had nothing to report.
    plant_misses = [
        (mid, r) for mid, r in sorted(plant_results.items()) if not r.ok
    ]
    behaviour_misses = [
        (mid, r) for mid, r in sorted(behaviour_results.items()) if not r.ok
    ]

    print()
    print(f"MISSES — plants ({len(plant_misses)}):")
    if not plant_misses:
        print("  none. every landmine is still in the corpus.")
    for mid, result in plant_misses:
        print(f"  {mid}  {result.detail}")

    print()
    print(f"MISSES — behaviour ({len(behaviour_misses)}):")
    if log is None:
        print(f"  not applicable: {log_error}")
    elif not behaviour_misses:
        print("  none. every observable expectation held.")
    for mid, result in behaviour_misses:
        print(f"  {mid}  {result.detail}")

    print()
    if plant_misses or behaviour_misses:
        return 1
    if log is None:
        print(
            "exit 2 — every plant is present, but no event log exists, so nothing\n"
            "about the system's behaviour has been observed. This is not a pass."
        )
        return 2
    return 0


if __name__ == "__main__":
    sys.exit(main(sys.argv[1:]))
